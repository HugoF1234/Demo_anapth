import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown_text: str, title: str) -> bytes:
    """Convert a markdown-formatted report into a .docx Word document."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Table detection
        if stripped.startswith("|") and i + 1 < len(lines):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # Title lines: **__TEXT__** or __TEXT__ (bold + underline, treated as heading)
        title_match = re.match(
            r"^\s*(?:\*\*)?__(.+?)__(?:\*\*)?\s*$", stripped
        )
        if title_match:
            text = title_match.group(1)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.underline = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # Markdown headings
        if stripped.startswith("# "):
            _add_heading(doc, stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            _add_heading(doc, stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            _add_heading(doc, stripped[4:].strip(), level=3)
        # Numbered list
        elif re.match(r"^\d+[\.\)]\s", stripped):
            content = re.sub(r"^\d+[\.\)]\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, content)
        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, stripped[2:])
        # Empty line
        elif stripped == "":
            pass
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, stripped)

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_heading(doc: Document, text: str, level: int):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.bold = True
    if level == 1:
        run.underline = True
        run.font.color.rgb = RGBColor(0, 0, 0)


def _add_rich_text(paragraph, text: str):
    """Parse inline markdown (bold, italic, bold+underline) into Word runs."""
    pattern = re.compile(
        r"(\*\*__.*?__\*\*"   # **__bold underline__**
        r"|\*\*\*.*?\*\*\*"   # ***bold italic***
        r"|\*\*.*?\*\*"       # **bold**
        r"|__.*?__"            # __underline bold__
        r"|\*.*?\*)"           # *italic*
    )
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**__") and part.endswith("__**"):
            run = paragraph.add_run(part[4:-4])
            run.bold = True
            run.underline = True
        elif part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.underline = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_table(doc: Document, table_lines: list[str]):
    rows_data = []
    for idx, line in enumerate(table_lines):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if idx == 1 and all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows_data.append(cells)

    if len(rows_data) < 1:
        return

    n_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row):
            if c_idx < n_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                _add_rich_text(p, cell_text)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
